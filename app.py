import streamlit as st
import time
import random
import PyPDF2
import docx
from io import BytesIO
from openai import OpenAI

# ==========================================
# 1. Page Configuration
# ==========================================
st.set_page_config(page_title="Multi-Agent Business Bureau", page_icon="👔", layout="wide")
st.title("🤖 Multi-Agent Business Analysis Bureau")
st.markdown("Upload documents and provide custom instructions. Our AI team will analyze the data exactly how you want it.")

# ==========================================
# 2. Engine Ignition
# ==========================================
DEEPSEEK_API_KEY = st.secrets["DEEPSEEK_API_KEY"]
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
MODEL_NAME = "deepseek-chat"

# ==========================================
# 3. Team Initialization (Agent Roles)
# ==========================================
agents = {
    "Intern": {
        "name": "Data Intern (Summarizer)",
        "icon": "🤓",
        "role_desc": "You are the Data Intern. Your job is to read the provided text and strictly summarize its core business information in under 800 words. Do not add any of your own analysis, just extract the facts."
    },
    "Researcher": {
        "name": "Chief Information Officer (Researcher)",
        "icon": "🕵️‍♂️",
        "role_desc": "You are the [Chief Information Officer]. Extract the core elements of the business model from the provided summaries: target audience, product/service, and revenue model. Output in a structured and concise format."
    },
    "Strategist": {
        "name": "Business Strategist (Strategist)",
        "icon": "🧠",
        "role_desc": "You are the [Business Strategist]. Based on the previous info, use frameworks like 'Porter's Five Forces' or 'SWOT Analysis' to deeply deconstruct the business logic, pointing out core competitiveness and market opportunities."
    },
    "Critic": {
        "name": "Critical Reviewer (Critic)",
        "icon": "🧐",
        "role_desc": "You are the [Critical Reviewer]. From a highly critical investor's perspective, identify fatal flaws, overly optimistic assumptions, and financial risks in this business model."
    },
    "Operations": {
        "name": "Operations Architect (Operations)",
        "icon": "🏗️",
        "role_desc": "You are the [Operations Architect]. Responsible for grounding the business ideas. You MUST use ASCII characters (like ├── and └──) to draw a clear [Organizational Chart Tree], showing what core departments need to be established. Then briefly explain how they collaborate in daily operations."
    },
    "Writer": {
        "name": "Chief Writer (Writer)",
        "icon": "✍️",
        "role_desc": "You are the [Chief Writer]. Summarize all the viewpoints of the above experts (especially the Operations Architect's tree chart), and write a well-structured 'Business Model Analysis Executive Summary'. It MUST include: 1. Project Overview 2. Advantage Analysis 3. Risks 4. Organizational Structure Tree 5. Final Investment Recommendation."
    }
}

# ==========================================
# 4. File Parsing, Compression & Word Generation Core 
# ==========================================
def read_uploaded_file(uploaded_file):
    extracted_text = ""
    try:
        if uploaded_file.name.endswith(".txt"):
            extracted_text = uploaded_file.getvalue().decode("utf-8")
        elif uploaded_file.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(uploaded_file)
            extracted_text = "".join(page.extract_text() for page in reader.pages if page.extract_text())
        elif uploaded_file.name.endswith(".docx"):
            doc = docx.Document(uploaded_file)
            extracted_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            
        if len(extracted_text) > 50000:
            extracted_text = extracted_text[:50000]
            
        return extracted_text
    except Exception as e:
        return f"[Error parsing {uploaded_file.name}]: {e}"

def summarize_long_text(text, filename):
    if len(text) < 3000:
        return text 
        
    messages = [
        {"role": "system", "content": agents["Intern"]["role_desc"] + "\nPlease output your summary in English."},
        {"role": "user", "content": f"Please summarize this document named '{filename}':\n\n{text}"}
    ]
    
    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=messages,
            temperature=0.3,
            max_tokens=1000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"[Summarization Failed for {filename}]: {e}"

def create_word_doc(report_text):
    doc = docx.Document()
    doc.add_heading('Multi-Agent Business Analysis Executive Summary', 0)
    doc.add_paragraph(report_text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 5. UI Layout & Main Logic
# ==========================================
with st.sidebar:
    st.header("📂 Data Input")
    uploaded_files = st.file_uploader(
        "Upload business documents", 
        type=["txt", "pdf", "docx"], 
        accept_multiple_files=True 
    )
    
    # 🚨 新增：自定义指令输入框
    st.header("🎯 Custom Instructions")
    custom_instruction = st.text_area(
        "Any specific focus? (Optional)", 
        placeholder="e.g., Focus heavily on the competitor analysis, or pay special attention to financial risks...",
        height=100
    )
    
    start_btn = st.button("🚀 Start Multi-Agent Analysis", type="primary", use_container_width=True)

if start_btn:
    if not uploaded_files:
        st.warning("👈 Please upload at least one file on the left sidebar!")
    else:
        st.subheader("📚 Phase 1: Data Digestion (Intern Agent)")
        all_compressed_context = ""
        
        for uploaded_file in uploaded_files:
            with st.status(f"Processing: {uploaded_file.name}...", expanded=True) as status:
                st.write("1. Extracting raw text...")
                raw_text = read_uploaded_file(uploaded_file)
                
                if raw_text.startswith("[Error"):
                    st.error(raw_text)
                    status.update(label=f"Failed: {uploaded_file.name}", state="error")
                    continue
                
                st.write(f"2. Generating summary...")
                compressed_text = summarize_long_text(raw_text, uploaded_file.name)
                all_compressed_context += f"\n\n--- SUMMARY OF {uploaded_file.name} ---\n{compressed_text}"
                
                st.write(f"✅ Summary generated.")
                status.update(label=f"Successfully processed {uploaded_file.name}", state="complete", expanded=False)

        if not all_compressed_context.strip():
            st.error("No valid text could be extracted from the uploaded files.")
            st.stop()

        # 🚨 新增：将用户的指令融入最初的 Prompt 中
        initial_prompt = f"Please read the following summaries of multiple business documents and conduct a deep analysis based on all of them combined:\n\n{all_compressed_context}"
        
        if custom_instruction.strip():
            initial_prompt += f"\n\n[USER'S SPECIAL INSTRUCTIONS]:\nPlease explicitly incorporate the following instructions into your thinking, analysis, and final output:\n{custom_instruction.strip()}"

        shared_context = [
            {"role": "user", "content": initial_prompt}
        ]
        
        st.divider()
        st.subheader("💬 Phase 2: Core Consulting Team Analysis")

        workflow_order = ["Researcher", "Strategist", "Critic", "Operations", "Writer"]
        
        for name in workflow_order:
            agent = agents[name]
            
            with st.chat_message(name, avatar=agent["icon"]):
                st.markdown(f"**{agent['name']}** is thinking...")
                
                thinking_time = random.uniform(2.0, 4.0)
                time.sleep(thinking_time)
                
                message_placeholder = st.empty()
                full_response = ""
                
                messages = [{"role": "system", "content": agent["role_desc"] + "\nPlease strictly adhere to your role and output your response entirely in English."}] + shared_context
                
                try:
                    stream = client.chat.completions.create(
                        model=MODEL_NAME,
                        messages=messages,
                        temperature=0.7,
                        max_tokens=2048,
                        stream=True 
                    )
                    
                    for chunk in stream:
                        if chunk.choices[0].delta.content is not None:
                            full_response += chunk.choices[0].delta.content
                            message_placeholder.markdown(full_response + "▌") 
                            
                    message_placeholder.markdown(full_response)
                    shared_context.append({"role": "assistant", "name": name, "content": f"[Output from {name}]:\n{full_response}"})
                    
                except Exception as e:
                    st.error(f"API Request Error: {e}")
                    st.stop()
                    
                time.sleep(1) 

        st.success("🎉 All agents have finished their analysis! The final executive summary is ready.")
        
        base_name = uploaded_files[0].name.split('.')[0] if uploaded_files else "Project"
        final_report_text = shared_context[-1]["content"]
        
        docx_buffer = create_word_doc(final_report_text)
        
        st.download_button(
            label="📥 Download Full Analysis Report (Word Doc)",
            data=docx_buffer,
            file_name=f"Report_Multi_{base_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
